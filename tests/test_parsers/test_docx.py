import tempfile
from docx import Document as DocxDocument
from doubase.parsers.docx import DocxParser


def _create_test_docx(path: str):
    doc = DocxDocument()
    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是一个包含内容的段落。")
    doc.add_paragraph("另一个段落。")
    doc.save(path)


def test_docx_supports():
    parser = DocxParser()
    assert parser.supports("/path/to/file.docx") is True
    assert parser.supports("/path/to/file.doc") is False
    assert parser.supports("/path/to/file.md") is False


def test_docx_parse_extracts_text():
    parser = DocxParser()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.close()
        _create_test_docx(f.name)
        doc = parser.parse(f.name)

    assert "测试文档" in doc.text
    assert "这是一个包含内容的段落" in doc.text
    assert "另一个段落" in doc.text
    assert doc.file_type == "docx"
    assert doc.source_path.endswith(".docx")
