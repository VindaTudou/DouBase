"""DOCX 文件解析器 — 将 .docx 转换为 Markdown 兼容文本。"""

from pathlib import Path

from doubase.parsers.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    """.docx 文件解析器。提取文本，将标题转换为 Markdown 格式。"""

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document as DocxDocument

        path = Path(file_path)
        doc = DocxDocument(str(path))

        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            if para.style.name.startswith("Heading"):
                level_str = para.style.name.split()[-1]
                level = int(level_str) if level_str.isdigit() else 1
                lines.append("#" * level + " " + text)
            else:
                lines.append(text)

        metadata = {}
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title

        return ParsedDocument(
            text="\n\n".join(lines),
            source_path=str(path.resolve()),
            file_type="docx",
            metadata=metadata,
        )
